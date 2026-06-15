"""
DOCX Report Generator — institutional research format.

References: Motilal Oswal (TCS/SBI) + Deven Choksey Research (TCS Q1FY26).

Layout:
  Cover page      2-col (key metrics left | CMP/TP/rating/thesis right)
                  Metrics strip (CMP | Target | Upside | MktCap | Rec | Sector)
  Page 2+         Result highlights + Quarterly snapshot
                  Key concall highlights (orange subheadings)
                  Annual performance + Scenario table
                  Estimate revision (new vs old)
                  Peer comparison (domestic + global)
                  Valuation (blended target + WACC)
                  Forensic accounting & quality
                  Financials — IS + BS (side by side)  |  CF + Ratios (side by side)
                  Key research findings
                  Rating history + Investment rating legend
                  Regulatory disclaimer
"""

from __future__ import annotations
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Dict, Any
from loguru import logger

from ..models.research import ResearchState, FindingType
from ..models.report import SectionType
from ..storage.storage_manager import StorageManager
from ..core.config import REPORT_CONFIG

# ── Colour palette ─────────────────────────────────────────────────────────
_NAV  = "1F497D"   # deep navy   — main headings
_TEAL = "17375E"   # dark blue   — primary section bars
_LBLU = "2F75B6"   # mid blue    — secondary section bars
_LBKG = "BDD7EE"   # pale blue   — table header fill
_ORG2 = "E26B0A"   # orange      — Deven Choksey topic subheadings / HOLD
_GRN  = "00B050"   # green       — BUY / positive
_RED  = "C00000"   # red         — SELL / negative
_ORG  = "FF9200"   # amber       — NEUTRAL / warning
_LGY  = "F2F2F2"   # light grey  — alternate table rows
_DGRY = "595959"   # dark grey   — secondary text
_WHT  = "FFFFFF"
_BLK  = "000000"

_RATING_HEX: dict[str, str] = {
    "BUY": _GRN, "STRONG BUY": _GRN, "OUTPERFORM": _GRN, "ACCUMULATE": _GRN,
    "HOLD": _ORG, "NEUTRAL": _ORG,
    "UNDERPERFORM": _RED, "REDUCE": _RED, "SELL": _RED, "STRONG SELL": _RED,
}
_NA = "N/A"


# ═══════════════════════════════════════════════════════════════════════════
# PUBLIC ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════

def generate_docx_report(state: ResearchState, storage: StorageManager) -> Optional[Path]:
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return None

    doc = Document()
    _set_margins(doc)

    profile     = state.company_profile or {}
    company     = state.company_name
    ticker      = state.ticker
    exchange    = profile.get("exchange", "")
    currency    = profile.get("currency", "INR")
    sector      = profile.get("sector", "")
    report_date = datetime.now().strftime("%d %B %Y")

    mktdata  = _payload(state, "04_market_data", "market_data", {})
    val      = state.valuation_summary or {}
    cmp_px   = mktdata.get("current_price")
    tgt_px   = state.target_price
    rating   = state.investment_rating or "UNDER REVIEW"
    risk_sc  = state.overall_risk_score

    upside = None
    if cmp_px and tgt_px and cmp_px > 0:
        upside = round((tgt_px - cmp_px) / cmp_px * 100, 1)

    fin_data = _collect_financial_data(state)

    # ── 1. Cover page ────────────────────────────────────────────────────
    _add_cover(doc, company, ticker, exchange, currency, sector,
               report_date, cmp_px, tgt_px, rating, upside, risk_sc,
               profile, mktdata, val, fin_data, state)

    doc.add_page_break()

    # ── 2. Page header + metrics strip ───────────────────────────────────
    _add_page_header(doc, company, ticker, report_date, sector)
    _add_metrics_strip(doc, currency, cmp_px, tgt_px, upside,
                       mktdata.get("market_cap"), rating, sector)

    # ── 3. Result highlights + quarterly snapshot ─────────────────────────
    _add_results_section(doc, state, company, ticker, report_date)
    _add_quarterly_snapshot(doc, state, currency)

    # ── 4. Key concall highlights (orange topic subheadings) ─────────────
    _add_concall_highlights(doc, state, company)

    # ── 5. Annual performance + scenario table ────────────────────────────
    _add_performance_table(doc, state, currency)

    # ── 6. Estimate revision table ────────────────────────────────────────
    _add_estimate_revision_table(doc, state, currency)

    # ── 7. Peer comparison ────────────────────────────────────────────────
    _add_peer_table(doc, state, company, currency)

    # ── 8. Valuation (blended target + WACC) ─────────────────────────────
    _add_valuation_section(doc, state, company, currency,
                           cmp_px, tgt_px, upside, rating, val)

    # ── 9. Forensic accounting ────────────────────────────────────────────
    _add_forensic_section(doc, state, company)

    # ── 10. Financials — side-by-side layout ──────────────────────────────
    _add_financials_section(doc, fin_data, currency, company)

    # ── 11. Key research findings ─────────────────────────────────────────
    _add_findings_section(doc, state)

    # ── 12. Analyst certification ─────────────────────────────────────────
    _add_certification(doc, company, ticker, report_date, state)

    # ── 13. Rating history + legend + disclaimer ──────────────────────────
    doc.add_page_break()
    _add_rating_history(doc, company, ticker)
    _add_rating_legend(doc)
    _add_disclaimer(doc, company, report_date)

    # ── Save ──────────────────────────────────────────────────────────────
    filename = f"{ticker}_{datetime.now().strftime('%Y%m%d_%H%M%S')}_equity_research.docx"
    out_path = storage.reports / filename
    doc.save(str(out_path))
    logger.info(f"DOCX report saved: {out_path}")
    return out_path


# ═══════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════

def _add_cover(doc, company, ticker, exchange, currency, sector,
               report_date, cmp_px, tgt_px, rating, upside, risk_sc,
               profile, mktdata, val, fin_data, state):
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # ── Header breadcrumb ─────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(
        f"India Equity Institutional Research   |   Initiating Coverage   |   {report_date}"
    )
    run.font.size = Pt(8)
    _color_run(run, _TEAL)

    # ── Company name ──────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(company.upper())
    run.bold = True
    run.font.size = Pt(18)
    _color_run(run, _NAV)

    # ── Report type tagline ───────────────────────────────────────────
    p = doc.add_paragraph()
    exec_s = state.report_sections.get(SectionType.EXECUTIVE_SUMMARY.value, "")
    headline = _extract_headline(exec_s, company)
    run = p.add_run(headline)
    run.bold = True
    run.font.size = Pt(11)
    _color_run(run, _ORG2)

    # ── Two-column layout ─────────────────────────────────────────────
    cover = doc.add_table(rows=1, cols=2)
    cover.style = "Table Grid"
    _remove_table_borders(cover)

    lw = Cm(7.2)
    rw = Cm(8.8)
    lc = cover.rows[0].cells[0]
    rc = cover.rows[0].cells[1]
    lc.width = lw
    rc.width = rw

    _left_column(lc, ticker, exchange, currency, profile, mktdata, fin_data, state, val)
    _right_column(rc, company, ticker, currency, cmp_px, tgt_px, rating,
                  upside, risk_sc, val, state, fin_data)

    # ── Analyst / platform line ───────────────────────────────────────
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(
        f"Equity Intelligence Agent (AI-assisted research)   |   "
        f"Platform: {REPORT_CONFIG.firm_name}   |   Run: {state.run_id}"
    )
    run.font.size = Pt(7)
    _color_run(run, _DGRY)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        "Investors are advised to refer through important disclosures at the last page of this report."
    )
    run2.font.size = Pt(7)
    run2.italic = True


def _left_column(cell, ticker, exchange, currency, profile, mktdata, fin_data, state, val):
    from docx.shared import Pt

    def _lp(text="", bold=False, size=8, color=_BLK):
        p = cell.add_paragraph()
        if text:
            r = p.add_run(text)
            r.bold = bold
            r.font.size = Pt(size)
            _color_run(r, color)
        return p

    _lp("Estimate change  ►", size=7, color=_DGRY)
    _lp("TP change        ►", size=7, color=_DGRY)
    _lp("Rating change    ►", size=7, color=_DGRY)
    _lp()

    shares = profile.get("shares_outstanding") or mktdata.get("shares_outstanding")
    mktcap = profile.get("market_cap") or mktdata.get("market_cap")
    hi52   = mktdata.get("week_52_high") or profile.get("52w_high")
    lo52   = mktdata.get("week_52_low")  or profile.get("52w_low")
    vol3m  = mktdata.get("avg_volume_3m")
    fv     = profile.get("face_value")
    beta   = mktdata.get("beta") or profile.get("beta")
    ff     = profile.get("free_float_pct") or mktdata.get("free_float_pct")

    metrics = [
        ("Bloomberg",           f"{ticker} IN" if exchange in ("NSE", "BSE") else ticker),
        ("Equity Shares (Mn)",  f"{shares/1e6:.0f}" if shares else _NA),
        (f"Mkt Cap ({currency} Mn)", f"{mktcap/1e6:,.0f}" if mktcap else _NA),
        ("52-Wk H / L",         f"{_fmt(hi52)} / {_fmt(lo52)}"),
        ("Vol. Avg (3m K)",      f"{vol3m/1e3:.0f}" if vol3m else _NA),
        ("Face Value",           f"{currency} {fv}" if fv else _NA),
        ("Beta",                 f"{beta:.2f}" if beta else _NA),
        ("Free Float (%)",       f"{ff:.1f}" if ff else _NA),
    ]
    _mini_table(cell, metrics, hdr_color=_LBLU, hdr_text="MARKET DATA")
    cell.add_paragraph()

    _add_fin_mini_table(cell, fin_data, currency)
    cell.add_paragraph()

    _add_shareholding_table(cell, state)
    cell.add_paragraph()

    # KPI callout boxes (Revenue CAGR + PAT CAGR) — Deven Choksey style
    _add_kpi_callouts(cell, state, fin_data)


def _right_column(cell, company, ticker, currency, cmp_px, tgt_px, rating,
                  upside, risk_sc, val, state, fin_data):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    rating_hex = _RATING_HEX.get((rating or "").upper(), _ORG)
    upside_str = f"({upside:+.1f}%)" if upside is not None else ""

    # CMP / TP / Rating banner
    banner_p = cell.add_paragraph()
    parts = [
        (f"CMP: {currency}{_fmt(cmp_px)}", _NAV, 11, False),
        (f"   TP: {currency}{_fmt(tgt_px)} {upside_str}", _NAV, 11, False),
        (f"   {rating.upper()}", rating_hex, 13, True),
    ]
    for text, color, size, bold in parts:
        r = banner_p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        _color_run(r, color)
    _shade_paragraph(banner_p, _LBKG)
    cell.add_paragraph()

    # Share price performance placeholder + market data box
    _add_cell_section_bar(cell, "SHARE PRICE PERFORMANCE", color=_TEAL)
    p = cell.add_paragraph()
    run = p.add_run(
        f"[12-month relative price performance vs. benchmark chart — "
        f"generated from market data agent output]"
    )
    run.font.size = Pt(8)
    run.italic = True
    _color_run(run, _DGRY)
    cell.add_paragraph()

    # Investment thesis bars
    _add_cell_section_bar(cell, "Investment Thesis", color=_LBLU)
    thesis = state.report_sections.get(SectionType.INVESTMENT_THESIS.value, "")
    _add_bullets_to_cell(cell, _extract_bullets(thesis, max_bullets=5), size=9)
    cell.add_paragraph()

    _add_cell_section_bar(cell, "Our View", color=_LBLU)
    risk_txt = state.report_sections.get(SectionType.RISK_ANALYSIS.value, "")
    _add_bullets_to_cell(cell, _extract_bullets(risk_txt, max_bullets=3), size=9)
    cell.add_paragraph()

    _add_cell_section_bar(cell, "Valuation & View", color=_LBLU)
    wacc   = val.get("wacc_inputs", {}).get("wacc", _NA)
    base   = _fmt(val.get("base_price"))
    bear   = _fmt(val.get("bear_price"))
    bull   = _fmt(val.get("bull_price"))
    val_bullets = [
        f"WACC: {wacc}%" if wacc != _NA else "WACC: N/A",
        f"Base target: {currency}{base}  |  Upside: {f'{upside:+.1f}%' if upside else _NA}",
        f"Bear: {currency}{bear}  |  Base: {currency}{base}  |  Bull: {currency}{bull}",
        f"Risk score: {risk_sc:.0f}/100" if risk_sc else "Risk score: N/A",
    ]
    _add_bullets_to_cell(cell, val_bullets, size=9)


# ═══════════════════════════════════════════════════════════════════════════
# PAGE HEADER & METRICS STRIP  (Deven Choksey style — appears top of page 2+)
# ═══════════════════════════════════════════════════════════════════════════

def _add_page_header(doc, company, ticker, report_date, sector):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    run = p.add_run(
        f"India Equity Institutional Research   ‖   Result Update   "
        f"‖   {report_date}   ‖   {ticker}"
    )
    run.font.size = Pt(8)
    _color_run(run, _DGRY)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p2 = doc.add_paragraph()
    run2 = p2.add_run(company)
    run2.bold = True
    run2.font.size = Pt(13)
    _color_run(run2, _NAV)


def _add_metrics_strip(doc, currency, cmp_px, tgt_px, upside, mktcap, rating, sector):
    """Single-row metrics bar: CMP | Target | Potential Upside | Mkt Cap | Recommendation | Sector."""
    from docx.shared import Pt
    rating_hex = _RATING_HEX.get((rating or "").upper(), _ORG)

    cols  = ["CMP*", "Target", "Potential Upside", f"Market Cap ({currency} Mn)", "Recommendation", "Sector"]
    vals  = [
        f"{currency} {_fmt(cmp_px)}",
        f"{currency} {_fmt(tgt_px)}",
        f"{f'{upside:+.1f}%' if upside is not None else _NA}",
        f"{mktcap/1e6:,.0f}" if mktcap else _NA,
        rating.upper() if rating else "UNDER REVIEW",
        sector or "Equity",
    ]

    t = doc.add_table(rows=2, cols=len(cols))
    t.style = "Table Grid"

    for ci, h in enumerate(cols):
        c = t.rows[0].cells[ci]
        c.text = h
        _set_cell_bg(c, _TEAL)
        if c.paragraphs[0].runs:
            r = c.paragraphs[0].runs[0]
            r.bold = True
            r.font.size = Pt(8)
            _color_run(r, _WHT)

    for ci, v in enumerate(vals):
        c = t.rows[1].cells[ci]
        c.text = v
        _set_cell_bg(c, _LBKG)
        if c.paragraphs[0].runs:
            r = c.paragraphs[0].runs[0]
            r.bold = True
            r.font.size = Pt(9)
            if ci == 4:
                _color_run(r, rating_hex)
            else:
                _color_run(r, _NAV)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════
# RESULT HIGHLIGHTS
# ═══════════════════════════════════════════════════════════════════════════

def _add_results_section(doc, state, company, ticker, report_date):
    _section_bar(doc, "Result Highlights")

    exec_text = state.report_sections.get(SectionType.EXECUTIVE_SUMMARY.value, "")
    if exec_text:
        _add_content_paragraphs(doc, exec_text, max_chars=2000)
    else:
        doc.add_paragraph(f"Executive summary pending data retrieval for {company}.")

    val_text = state.report_sections.get(SectionType.VALUATION.value, "")
    if val_text:
        _section_bar(doc, "Outlook & Valuation", color=_LBLU)
        _add_content_paragraphs(doc, val_text, max_chars=800)


def _add_quarterly_snapshot(doc, state, currency):
    """Annual snapshot (last 2 fiscal years) with YoY — agent 03 only produces annual data."""
    _section_bar(doc, "Result Snapshot (Annual YoY)")

    history = state.financial_history or {}
    annual_years = sorted([k for k in history.keys() if len(k) == 4])
    if len(annual_years) < 2:
        doc.add_paragraph("Insufficient annual data for comparison.")
        return

    cur  = annual_years[-1]
    prv  = annual_years[-2]
    yago = annual_years[-3] if len(annual_years) >= 3 else annual_years[0]

    def _qv(yr, *keys):
        d = history.get(yr, {})
        for sect in ("income_statements", "balance_sheets", "cash_flows"):
            sd = d.get(sect) or {}
            for k in keys:
                v = sd.get(k)
                if v is not None:
                    return v
        return None

    def _chg(v1, v2):
        try:
            return f"{(v1 - v2) / abs(v2) * 100:.1f}%"
        except Exception:
            return _NA

    rows_def = [
        ("Revenue",           ["revenue"]),
        ("Direct Expenses",   ["direct_expense", "cogs"]),
        ("Other Expenses",    ["other_expenses", "opex_other"]),
        ("Total Expenses",    ["total_expenses", "total_opex"]),
        ("**EBITDA",          ["ebitda"]),
        ("  EBITDA Margin %", ["ebitda_margin_pct", "ebitda_margin"]),
        ("Depreciation",      ["depreciation", "da"]),
        ("**EBIT",            ["ebit"]),
        ("  EBIT Margin %",   ["ebit_margin_pct", "ebit_margin"]),
        ("Finance Costs",     ["finance_costs", "interest_expense"]),
        ("Other Income",      ["other_income"]),
        ("Pretax Income",     ["pbt", "pre_tax_income"]),
        ("Tax Expense",       ["tax_expense", "income_tax"]),
        ("Net Profit",        ["net_income", "pat"]),
        ("Minority Interest", ["minority_interest"]),
        ("**PAT (after MI)",  ["pat_after_mi", "pat"]),
        ("  Net Margin %",    ["net_margin_pct", "pat_margin_pct"]),
        ("**Diluted EPS",     ["diluted_eps", "eps"]),
    ]

    hdr = [f"Particulars ({currency} Mn)", cur, prv, yago, "YoY (cur vs prv)", "2Y (cur vs yago)"]
    t = doc.add_table(rows=1 + len(rows_def), cols=len(hdr))
    t.style = "Table Grid"

    for ci, h in enumerate(hdr):
        c = t.rows[0].cells[ci]
        c.text = h
        _set_cell_bg(c, _ORG2)
        if c.paragraphs[0].runs:
            r = c.paragraphs[0].runs[0]
            r.bold = True
            r.font.size = _pt(8)
            _color_run(r, _WHT)

    for ri, (label, keys) in enumerate(rows_def):
        is_bold = label.startswith("**") or "Margin" in label
        clean   = label.lstrip("*").strip()
        vc  = _qv(cur,  *keys)
        vp  = _qv(prv,  *keys)
        vya = _qv(yago, *keys)

        bg = _LBKG if label.startswith("**") else (_LGY if ri % 2 else _WHT)
        vals = [clean, _fmt(vc), _fmt(vp), _fmt(vya),
                _chg(vc, vp) if vc and vp else _NA,
                _chg(vc, vya) if vc and vya else _NA]

        for ci, val in enumerate(vals):
            c = t.rows[ri + 1].cells[ci]
            c.text = str(val)
            _set_cell_bg(c, bg)
            if c.paragraphs[0].runs:
                r = c.paragraphs[0].runs[0]
                r.font.size = _pt(8)
                r.bold = is_bold and ci == 0

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════
# CONCALL HIGHLIGHTS  (Deven Choksey orange topic subheadings)
# ═══════════════════════════════════════════════════════════════════════════

def _add_concall_highlights(doc, state, company):
    from docx.shared import Pt

    _section_bar(doc, "Key Concall Highlights")

    # Agent 12 produces: guidance_summary (dict), sentiment (str), llm_transcript_analysis (str)
    tr_out      = state.agent_outputs.get("12_transcript_retrieval")
    tr_pay      = (tr_out.payload if tr_out else {}) or {}
    sentiment   = tr_pay.get("sentiment", "")
    mgmt_tone   = ""   # not produced by agent 12; omit
    guidance_raw = tr_pay.get("guidance_summary", {})
    llm_analysis = tr_pay.get("llm_transcript_analysis", "")

    # Convert guidance_summary dict to bullet list
    if isinstance(guidance_raw, dict):
        guidance_bullets = [f"{k}: {v}" for k, v in guidance_raw.items() if v]
    elif isinstance(guidance_raw, list):
        guidance_bullets = [str(x) for x in guidance_raw]
    else:
        guidance_bullets = []

    # Extract bullets from LLM analysis as fallback content
    llm_bullets = _extract_bullets(llm_analysis, max_bullets=12) if llm_analysis else []

    # Group available content into topic buckets
    topics = {
        "Management Guidance": guidance_bullets[:4] if guidance_bullets else llm_bullets[:3],
        "Revenue and Demand Outlook": llm_bullets[3:6] if len(llm_bullets) > 3 else [],
        "Margin Commentary": llm_bullets[6:9] if len(llm_bullets) > 6 else [],
        "Deal Wins / Strategic Initiatives": [],
        "Outlook and Key Monitorables": llm_bullets[9:12] if len(llm_bullets) > 9 else [],
    }

    has_content = False
    for topic, bullets in topics.items():
        if not bullets:
            continue
        has_content = True
        # Orange topic subheading
        p = doc.add_paragraph()
        run = p.add_run(topic)
        run.bold = True
        run.font.size = Pt(10)
        _color_run(run, _ORG2)

        for b in (bullets if isinstance(bullets, list) else [str(bullets)])[:5]:
            p2 = doc.add_paragraph(str(b), style="List Bullet")
            if p2.runs:
                p2.runs[0].font.size = Pt(9)

    if not has_content:
        # Fall back to investment thesis sections
        for topic, section_key in [
            ("Revenue and Profitability", SectionType.EXECUTIVE_SUMMARY.value),
            ("Investment Thesis and Drivers", SectionType.INVESTMENT_THESIS.value),
            ("Risk Factors", SectionType.RISK_ANALYSIS.value),
        ]:
            text = state.report_sections.get(section_key, "")
            if not text:
                continue
            p = doc.add_paragraph()
            run = p.add_run(topic)
            run.bold = True
            run.font.size = Pt(10)
            _color_run(run, _ORG2)
            for b in _extract_bullets(text, max_bullets=4):
                p2 = doc.add_paragraph(b, style="List Bullet")
                if p2.runs:
                    p2.runs[0].font.size = Pt(9)

    if sentiment:
        p = doc.add_paragraph()
        run = p.add_run(f"Management Sentiment: {sentiment}")
        run.font.size = Pt(9)
        run.italic = True
        _color_run(run, _DGRY)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════
# ANNUAL PERFORMANCE TABLE
# ═══════════════════════════════════════════════════════════════════════════

def _add_performance_table(doc, state, currency):
    from docx.shared import Pt

    doc.add_paragraph()
    _section_bar(doc, "Annual Performance Summary")

    fm_out    = state.agent_outputs.get("07_financial_modeling")
    forecasts = (fm_out.payload.get("forecasts", {}) if fm_out else {})
    base_f    = forecasts.get("BASE", [])
    history   = state.financial_history or {}
    years_hist = sorted([k for k in history.keys() if len(k) == 4])[-4:]

    est_years = [str(f.get("year", f"FY+{i+1}E")) for i, f in enumerate(base_f[:3])]
    all_cols  = years_hist + est_years
    if not all_cols:
        doc.add_paragraph("Performance data not yet available.")
        return

    hdr = ["Metric"] + all_cols
    rows: List[List[str]] = []

    def _hval(yr, *keys):
        d = history.get(yr, {})
        for sect in ("income_statements", "balance_sheets", "cash_flows"):
            sd = d.get(sect) or {}
            for k in keys:
                v = sd.get(k)
                if v is not None:
                    return v
        return None

    def _fval(i, key):
        if i < len(base_f):
            return base_f[i].get(key)
        return None

    metrics = [
        ("Revenue",          ["revenue"],              "revenue"),
        ("YoY Growth %",     ["revenue_growth_pct"],   "revenue_growth_pct"),
        ("**EBITDA",         ["ebitda"],               "ebitda"),
        ("  EBITDA Mgn %",   ["ebitda_margin_pct"],    "ebitda_margin"),
        ("**EBIT",           ["ebit"],                 "ebit"),
        ("  EBIT Mgn %",     ["ebit_margin_pct"],      "ebit_margin"),
        ("**PAT",            ["net_income", "pat"],    "net_income"),
        ("  PAT Mgn %",      ["net_margin_pct"],       "net_margin"),
        ("**EPS",            ["eps"],                  "eps"),
        ("Free Cash Flow",   ["fcf", "free_cash_flow"], "fcf"),
        ("RoE %",            ["roe_pct"],              ""),
        ("RoCE %",           ["roce_pct"],             ""),
    ]

    for label, hkeys, fkey in metrics:
        row = [label]
        for yr in years_hist:
            v = _hval(yr, *hkeys)
            row.append(_fmt_metric(v, label))
        for i in range(len(est_years)):
            v = _fval(i, fkey)
            row.append((_fmt_metric(v, label) + "E") if v is not None else _NA)
        rows.append(row)

    t = doc.add_table(rows=1 + len(rows), cols=len(hdr))
    t.style = "Table Grid"

    for ci, h in enumerate(hdr):
        c = t.rows[0].cells[ci]
        c.text = h
        _set_cell_bg(c, _TEAL)
        if c.paragraphs[0].runs:
            r = c.paragraphs[0].runs[0]
            r.bold = True
            r.font.size = _pt(8)
            _color_run(r, _WHT)

    for ri, row in enumerate(rows):
        is_bold = row[0].startswith("**")
        bg = _LBKG if is_bold else (_LGY if ri % 2 else _WHT)
        for ci, v in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = str(v).lstrip("*")
            _set_cell_bg(c, bg)
            if c.paragraphs[0].runs:
                r = c.paragraphs[0].runs[0]
                r.font.size = _pt(8)
                r.bold = is_bold

    doc.add_paragraph()
    _section_bar(doc, "Scenario Analysis (Bear / Base / Bull)", color=_LBLU)
    _add_scenario_table(doc, state, currency)


def _add_scenario_table(doc, state, currency):
    scen_out  = state.agent_outputs.get("15_scenario_analysis")
    scen_pay  = scen_out.payload if scen_out else {}
    fm_out    = state.agent_outputs.get("07_financial_modeling")
    forecasts = (fm_out.payload.get("forecasts", {}) if fm_out else {})

    labels = ["BEAR", "BASE", "BULL"]
    hdr = ["Metric"] + [s.title() + " Case" for s in labels]
    rows_data = [
        ("Revenue CAGR (5Y %)", "revenue_cagr_pct", "%"),
        ("EBITDA Margin (%)",   "ebitda_margin",    "%"),
        ("PAT CAGR (5Y %)",     "pat_cagr_pct",     "%"),
        ("Target Price",        "implied_price",    currency),   # scenario agent uses implied_price
        ("Upside/Downside",     "upside_pct",       "%"),
        ("Probability",         "_probability",     "%"),        # comes from probability_weights dict
    ]

    # Agent 15 stores scenarios under scen_pay["scenarios"]["BEAR/BASE/BULL"]
    scenarios_dict    = scen_pay.get("scenarios", {}) or {}
    probability_weights = scen_pay.get("probability_weights", {}) or {}

    t = doc.add_table(rows=1 + len(rows_data), cols=len(hdr))
    t.style = "Table Grid"
    for ci, h in enumerate(hdr):
        c = t.rows[0].cells[ci]
        c.text = h
        _set_cell_bg(c, _TEAL)
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].runs[0].font.size = _pt(9)
            _color_run(c.paragraphs[0].runs[0], _WHT)

    for ri, (lbl, key, unit) in enumerate(rows_data):
        t.rows[ri + 1].cells[0].text = lbl
        for ci, s in enumerate(labels):
            cell = t.rows[ri + 1].cells[ci + 1]
            if key == "_probability":
                raw = probability_weights.get(s)
                v = raw * 100 if isinstance(raw, float) and raw <= 1 else raw
            else:
                sd = scenarios_dict.get(s, {}) or (forecasts.get(s, [{}]) or [{}])[0]
                v  = sd.get(key)
            if v is not None:
                cell.text = f"{unit}{v:.1f}" if unit == currency else f"{v:.1f}{unit}"
            else:
                cell.text = _NA
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = _pt(9)

    pwtp = scen_pay.get("probability_weighted_target")  # was wrongly "probability_weighted_target_price"
    if pwtp:
        p = doc.add_paragraph()
        r = p.add_run(f"Probability-Weighted Target Price: {currency}{pwtp:.2f}")
        r.bold = True
        r.font.size = _pt(10)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════
# ESTIMATE REVISION TABLE  (Deven Choksey — New vs Old vs Variation)
# ═══════════════════════════════════════════════════════════════════════════

def _add_estimate_revision_table(doc, state, currency):
    _section_bar(doc, "Change in Estimates")

    fm_out = state.agent_outputs.get("07_financial_modeling")
    if not fm_out:
        doc.add_paragraph("Financial modelling data not available.")
        return

    new_f = fm_out.payload.get("forecasts", {}).get("BASE", [])[:3]
    old_f = fm_out.payload.get("prior_forecasts", {}).get("BASE", [])[:3] if fm_out.payload.get("prior_forecasts") else []

    if not new_f:
        doc.add_paragraph("Forecast data not available.")
        return

    # Agent 07 produces "llm_modeling_commentary", not "revision_commentary"
    revision_text = fm_out.payload.get("llm_modeling_commentary", "")
    if revision_text:
        _add_content_paragraphs(doc, revision_text, max_chars=600)
    else:
        doc.add_paragraph(
            "Estimates have been updated to reflect the latest quarterly performance and "
            "management commentary. Refer to the scenario section for Bear/Base/Bull ranges."
        ).runs[0].font.size = _pt(9)

    new_yrs = [str(f.get("year", f"FY+{i+1}E")) for i, f in enumerate(new_f)]
    old_yrs = [str(f.get("year", f"FY+{i+1}E")) for i, f in enumerate(old_f)] if old_f else []

    if old_f:
        # Three-block table: New | Old | Variation
        hdr_cells = ["Metric"] + [f"New {y}E" for y in new_yrs] + [f"Old {y}E" for y in old_yrs] + [f"Var {y}E" for y in new_yrs]
        rows_def = [
            ("Revenue",   "revenue"),
            ("EBIT",      "ebit"),
            ("PAT",       "net_income"),
            ("EPS",       "eps"),
            ("EBIT %",    "ebit_margin"),
            ("PAT %",     "net_margin"),
        ]
        t = doc.add_table(rows=1 + len(rows_def), cols=len(hdr_cells))
        t.style = "Table Grid"

        # Header with colour blocks
        for ci, h in enumerate(hdr_cells):
            c = t.rows[0].cells[ci]
            c.text = h
            bg = _TEAL if ci == 0 else (_GRN if ci <= len(new_yrs) else (_RED if ci <= len(new_yrs) + len(old_yrs) else _LBLU))
            _set_cell_bg(c, bg)
            if c.paragraphs[0].runs:
                r = c.paragraphs[0].runs[0]
                r.bold = True
                r.font.size = _pt(8)
                _color_run(r, _WHT)

        for ri, (lbl, key) in enumerate(rows_def):
            t.rows[ri + 1].cells[0].text = lbl
            n_vals = [f.get(key) for f in new_f]
            o_vals = [f.get(key) for f in old_f] if old_f else [None] * len(new_f)
            for i, v in enumerate(n_vals):
                t.rows[ri + 1].cells[1 + i].text = _fmt(v) if v else _NA
            for i, v in enumerate(o_vals):
                t.rows[ri + 1].cells[1 + len(new_yrs) + i].text = _fmt(v) if v else _NA
            for i, (nv, ov) in enumerate(zip(n_vals, o_vals)):
                try:
                    pct = (nv - ov) / abs(ov) * 100
                    t.rows[ri + 1].cells[1 + len(new_yrs) + len(old_yrs) + i].text = f"{pct:+.1f}%"
                except Exception:
                    t.rows[ri + 1].cells[1 + len(new_yrs) + len(old_yrs) + i].text = _NA
            for ci in range(len(hdr_cells)):
                if t.rows[ri + 1].cells[ci].paragraphs[0].runs:
                    t.rows[ri + 1].cells[ci].paragraphs[0].runs[0].font.size = _pt(8)
    else:
        # Only new estimates
        hdr_cells = ["Metric"] + [f"{y}E" for y in new_yrs]
        rows_def = [("Revenue", "revenue"), ("EBIT", "ebit"), ("PAT", "net_income"),
                    ("EPS", "eps"), ("EBIT %", "ebit_margin"), ("PAT %", "net_margin")]
        t = doc.add_table(rows=1 + len(rows_def), cols=len(hdr_cells))
        t.style = "Table Grid"
        for ci, h in enumerate(hdr_cells):
            c = t.rows[0].cells[ci]
            c.text = h
            _set_cell_bg(c, _TEAL)
            if c.paragraphs[0].runs:
                r = c.paragraphs[0].runs[0]
                r.bold = True; r.font.size = _pt(8); _color_run(r, _WHT)
        for ri, (lbl, key) in enumerate(rows_def):
            t.rows[ri + 1].cells[0].text = lbl
            for i, f in enumerate(new_f):
                t.rows[ri + 1].cells[i + 1].text = _fmt(f.get(key))
            for ci in range(len(hdr_cells)):
                if t.rows[ri + 1].cells[ci].paragraphs[0].runs:
                    t.rows[ri + 1].cells[ci].paragraphs[0].runs[0].font.size = _pt(8)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════
# PEER COMPARISON TABLE  (Deven Choksey style — domestic + global)
# ═══════════════════════════════════════════════════════════════════════════

def _add_peer_table(doc, state, company, currency):
    _section_bar(doc, "Peer Comparison")

    # Peer data lives in agent 04 (market data), not a separate peer agent
    mkt_out   = state.agent_outputs.get("04_market_data")
    mkt_data  = (mkt_out.payload.get("market_data", {}) if mkt_out else {}) or {}
    peers_raw = mkt_data.get("peer_market_data", []) or []

    valuation_narrative = ""  # no dedicated peer valuation narrative available
    if valuation_narrative:
        _add_content_paragraphs(doc, valuation_narrative, max_chars=500)

    hdr = [
        "Company", f"CMP ({currency})", "Mkt Cap (USD Bn)",
        "Rev CAGR", "EPS CAGR", "EBIT Mgn",
        "P/E FY1E", "P/E FY2E",
        "FCF Yield FY1", "FCF Yield FY2",
        "RoE FY1", "RoE FY2",
    ]

    # Subject company row from state
    val    = state.valuation_summary or {}
    mktd   = _payload(state, "04_market_data", "market_data", {})
    cmp_px = mktd.get("current_price")
    mktcap = mktd.get("market_cap")
    mktcap_usd = mktcap / 1e9 / 83 if mktcap else None   # rough INR→USD

    fm_out = state.agent_outputs.get("07_financial_modeling")
    base_f = (fm_out.payload.get("forecasts", {}).get("BASE", []) if fm_out else [])

    def _cagr(fs, key, years=2):
        try:
            v0 = base_f[0].get(key)
            v1 = base_f[min(years, len(base_f) - 1)].get(key)
            return ((v1 / v0) ** (1 / years) - 1) * 100 if v0 and v1 else None
        except Exception:
            return None

    subject_row = [
        company[:20],
        _fmt(cmp_px) if cmp_px else _NA,
        f"{mktcap_usd:.0f}" if mktcap_usd else _NA,
        f"{_cagr(base_f, 'revenue'):.1f}%" if _cagr(base_f, 'revenue') else _NA,
        f"{_cagr(base_f, 'eps'):.1f}%" if _cagr(base_f, 'eps') else _NA,
        _NA, _NA, _NA, _NA, _NA, _NA, _NA,
    ]

    def _peer_row(p):
        return [
            str(p.get("name", p.get("ticker", "?")))[:20],
            _fmt(p.get("cmp") or p.get("current_price")),
            _fmt(p.get("market_cap_usd_bn")),
            f"{p.get('revenue_cagr_pct', _NA):.1f}%" if isinstance(p.get("revenue_cagr_pct"), (int, float)) else _NA,
            f"{p.get('eps_cagr_pct', _NA):.1f}%" if isinstance(p.get("eps_cagr_pct"), (int, float)) else _NA,
            f"{p.get('ebit_margin_pct', _NA):.1f}%" if isinstance(p.get("ebit_margin_pct"), (int, float)) else _NA,
            f"{p.get('pe_fy1', _NA):.1f}x" if isinstance(p.get("pe_fy1"), (int, float)) else _NA,
            f"{p.get('pe_fy2', _NA):.1f}x" if isinstance(p.get("pe_fy2"), (int, float)) else _NA,
            f"{p.get('fcf_yield_fy1', _NA):.1f}%" if isinstance(p.get("fcf_yield_fy1"), (int, float)) else _NA,
            f"{p.get('fcf_yield_fy2', _NA):.1f}%" if isinstance(p.get("fcf_yield_fy2"), (int, float)) else _NA,
            f"{p.get('roe_fy1', _NA):.1f}%" if isinstance(p.get("roe_fy1"), (int, float)) else _NA,
            f"{p.get('roe_fy2', _NA):.1f}%" if isinstance(p.get("roe_fy2"), (int, float)) else _NA,
        ]

    domestic = [p for p in peers_raw if p.get("peer_type", "domestic").lower() in ("domestic", "local", "india")]
    global_p = [p for p in peers_raw if p.get("peer_type", "domestic").lower() in ("global", "international")]
    if not domestic and not global_p:
        domestic = peers_raw

    all_rows: List[tuple[str, list]] = [("subject", subject_row)]
    if domestic:
        all_rows += [("dom", _peer_row(p)) for p in domestic[:6]]
        # Mean/Median for domestic peers (skip for now, would require per-column aggregation)
    if global_p:
        all_rows += [("global", _peer_row(p)) for p in global_p[:4]]

    t = doc.add_table(rows=1 + len(all_rows), cols=len(hdr))
    t.style = "Table Grid"

    for ci, h in enumerate(hdr):
        c = t.rows[0].cells[ci]
        c.text = h
        _set_cell_bg(c, _TEAL)
        if c.paragraphs[0].runs:
            r = c.paragraphs[0].runs[0]
            r.bold = True; r.font.size = _pt(7); _color_run(r, _WHT)

    section_labels = {0: "Domestic Peers", len(domestic) + 1 if domestic else 9999: "Global Peers"}

    ri_t = 1
    for kind, row_data in all_rows:
        if kind == "subject":
            bg = _LBKG
        elif kind == "dom":
            bg = _LGY
        else:
            bg = _WHT

        for ci, v in enumerate(row_data):
            c = t.rows[ri_t].cells[ci]
            c.text = str(v)
            _set_cell_bg(c, bg)
            if c.paragraphs[0].runs:
                r = c.paragraphs[0].runs[0]
                r.font.size = _pt(7)
                r.bold = (kind == "subject")
        ri_t += 1

    doc.add_paragraph()

    # Source citation
    p = doc.add_paragraph()
    run = p.add_run("Source: Company, Bloomberg, Platform estimates")
    run.font.size = _pt(8)
    run.italic = True
    _color_run(run, _DGRY)
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════
# VALUATION
# ═══════════════════════════════════════════════════════════════════════════

def _add_valuation_section(doc, state, company, currency, cmp_px, tgt_px, upside, rating, val):
    _section_bar(doc, "Valuation")

    wacc_in = val.get("wacc_inputs", {}) or {}
    wacc_val = wacc_in.get("wacc", _NA)

    # Valuation narrative — agent 08 stores as "llm_valuation_commentary"
    val_narr = _payload(state, "08_valuation", "llm_valuation_commentary", "") or \
               state.report_sections.get(SectionType.VALUATION.value, "")
    if val_narr:
        _add_content_paragraphs(doc, val_narr, max_chars=800)

    # Extract per-method implied values from nested ValuationSummary structure
    base_case  = val.get("base_case", {}) or {}
    dcf_val    = base_case.get("dcf", {}).get("intrinsic_value_per_share") if base_case.get("dcf") else None
    rel_val    = base_case.get("relative", {}).get("blended_value_per_share") if base_case.get("relative") else None
    sotp_val   = base_case.get("sotp", {}).get("intrinsic_value_per_share") if base_case.get("sotp") else None

    # Blended target table
    rows = [
        ("Valuation Method",     "Weight",   "Implied Value"),
        ("DCF (Base Case)",      "50%",      f"{currency}{_fmt(dcf_val)}" if dcf_val else _NA),
        ("Peer Relative / EV",   "50%",      f"{currency}{_fmt(rel_val)}" if rel_val else _NA),
        ("SOTP",                 "—",        f"{currency}{_fmt(sotp_val)}" if sotp_val else _NA),
        ("Blended Target Price", "100%",     f"{currency}{_fmt(tgt_px)}"),
        ("Current Market Price", "—",        f"{currency}{_fmt(cmp_px)}"),
        ("Upside / (Downside)",  "—",        f"{f'{upside:+.1f}%' if upside else _NA}"),
    ]

    t = doc.add_table(rows=len(rows), cols=3)
    t.style = "Table Grid"
    for ri, row in enumerate(rows):
        bg = _TEAL if ri == 0 else (_LBKG if ri >= len(rows) - 3 else (_LGY if ri % 2 else _WHT))
        for ci, txt in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.text = txt
            _set_cell_bg(c, bg)
            if c.paragraphs[0].runs:
                r = c.paragraphs[0].runs[0]
                r.font.size = _pt(9)
                r.bold = (ri == 0 or ri >= len(rows) - 3)
                if ri == 0:
                    _color_run(r, _WHT)

    doc.add_paragraph()

    # WACC construction
    if wacc_in:
        _section_bar(doc, "WACC Construction", color=_LBLU)
        wacc_rows = [
            ("Risk-free Rate",         f"{_fmt(wacc_in.get('risk_free_rate'))}%"),
            ("Equity Risk Premium",    f"{_fmt(wacc_in.get('equity_risk_premium'))}%"),
            ("Country Risk Premium",   f"{_fmt(wacc_in.get('country_risk_premium'))}%"),
            ("Beta",                   f"{_fmt(wacc_in.get('beta'))}"),
            ("Cost of Equity",         f"{_fmt(wacc_in.get('cost_of_equity'))}%"),
            ("Pre-tax Cost of Debt",   f"{_fmt(wacc_in.get('pre_tax_cost_of_debt'))}%"),
            ("Effective Tax Rate",     f"{_fmt(wacc_in.get('tax_rate'))}%"),
            ("After-tax Cost of Debt", f"{_fmt(wacc_in.get('after_tax_cost_of_debt'))}%"),
            ("Equity Weight",          f"{_fmt(wacc_in.get('equity_weight'), pct=True)}%"),
            ("Debt Weight",            f"{_fmt(wacc_in.get('debt_weight'), pct=True)}%"),
            ("WACC",                   f"{_fmt(wacc_in.get('wacc'))}%"),
        ]
        t2 = doc.add_table(rows=len(wacc_rows), cols=2)
        t2.style = "Table Grid"
        for ri, (k, v) in enumerate(wacc_rows):
            is_last = ri == len(wacc_rows) - 1
            bg = _TEAL if is_last else (_LGY if ri % 2 else _WHT)
            for ci, txt in enumerate((k, v)):
                c = t2.rows[ri].cells[ci]
                c.text = txt
                _set_cell_bg(c, bg)
                if c.paragraphs[0].runs:
                    r = c.paragraphs[0].runs[0]
                    r.font.size = _pt(9)
                    r.bold = is_last
                    if is_last:
                        _color_run(r, _WHT)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════
# FORENSIC ACCOUNTING
# ═══════════════════════════════════════════════════════════════════════════

def _add_forensic_section(doc, state, company):
    _section_bar(doc, "Forensic Accounting & Quality Assessment")

    forensic = _payload(state, "06_forensic_accounting", "details", {})
    # Agent 05 exposes scores as direct payload keys, not nested in "quality_metrics"
    cash_conv_score = _payload(state, "05_accounting_quality", "cash_conversion_score", _NA)
    accrual_score   = _payload(state, "05_accounting_quality", "accrual_score", _NA)
    rev_qual_score  = _payload(state, "05_accounting_quality", "revenue_quality_score", _NA)

    scores_data = [
        ("Score", "Value", "Threshold", "Interpretation"),
        ("Beneish M-Score",
         str(forensic.get("beneish", {}).get("m_score", _NA)),
         "< −1.78 (no manipulation)",
         forensic.get("beneish", {}).get("classification", _NA)),
        ("Piotroski F-Score",
         f"{forensic.get('piotroski', {}).get('f_score', _NA)} / 9",
         "≥ 7 = strong, ≤ 2 = weak",
         forensic.get("piotroski", {}).get("interpretation", _NA)),
        ("Altman Z-Score",
         str(forensic.get("altman", {}).get("z_score", _NA)),
         "> 2.6 = safe zone",
         forensic.get("altman", {}).get("zone", _NA)),
        ("Sloan Accrual Ratio",
         str(forensic.get("sloan_accrual", _NA)),
         "< 5% preferred",
         ""),
        ("Cash Conversion (OCF/NI)",
         str(cash_conv_score),
         "> 0.8× preferred",
         ""),
        ("Accrual Score",
         str(accrual_score),
         "< 0.1 preferred",
         ""),
        ("Revenue Quality Score",
         str(rev_qual_score),
         "> 0.7 preferred",
         ""),
    ]

    t = doc.add_table(rows=len(scores_data), cols=4)
    t.style = "Table Grid"
    for ri, row in enumerate(scores_data):
        bg = _TEAL if ri == 0 else (_LGY if ri % 2 else _WHT)
        for ci, txt in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.text = str(txt)
            _set_cell_bg(c, bg)
            if c.paragraphs[0].runs:
                r = c.paragraphs[0].runs[0]
                r.font.size = _pt(9)
                if ri == 0:
                    r.bold = True; _color_run(r, _WHT)

    doc.add_paragraph()

    red_flags = [f for f in state.all_findings if f.finding_type == FindingType.RED_FLAG]
    if red_flags:
        _section_bar(doc, "Forensic Red Flags", color=_RED)
        t2 = doc.add_table(rows=1 + len(red_flags[:12]), cols=4)
        t2.style = "Table Grid"
        for ci, h in enumerate(("#", "Flag", "Severity", "Implication")):
            c = t2.rows[0].cells[ci]
            c.text = h
            _set_cell_bg(c, _TEAL)
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].bold = True
                c.paragraphs[0].runs[0].font.size = _pt(9)
                _color_run(c.paragraphs[0].runs[0], _WHT)
        for ri, f in enumerate(red_flags[:12]):
            t2.rows[ri + 1].cells[0].text = str(ri + 1)
            t2.rows[ri + 1].cells[1].text = f.title[:80]
            t2.rows[ri + 1].cells[2].text = f.risk_level.value
            t2.rows[ri + 1].cells[3].text = (f.evidence or "")[:120]
            for ci in range(4):
                if t2.rows[ri + 1].cells[ci].paragraphs[0].runs:
                    t2.rows[ri + 1].cells[ci].paragraphs[0].runs[0].font.size = _pt(8)

    doc.add_paragraph()
    aq_text = state.report_sections.get(SectionType.ACCOUNTING_QUALITY.value, "")
    if aq_text:
        _section_bar(doc, "Accounting Quality Assessment", color=_LBLU)
        _add_content_paragraphs(doc, aq_text, max_chars=1200)


# ═══════════════════════════════════════════════════════════════════════════
# FINANCIALS — Side-by-side layout (Deven Choksey Exhibit 1/2/3/4 style)
# ═══════════════════════════════════════════════════════════════════════════

def _add_financials_section(doc, fin_data: dict, currency: str, company: str):
    doc.add_page_break()
    _section_bar(doc, f"Key Financials   ({currency} Mn)")

    years    = fin_data.get("years", [])
    is_rows  = fin_data.get("income_statement", [])
    bs_rows  = fin_data.get("balance_sheet", [])
    cf_rows  = fin_data.get("cash_flow", [])
    rat_rows = fin_data.get("ratios", [])

    if not years:
        doc.add_paragraph("Financial data not available.")
        return

    # ── P&L and BS side by side ───────────────────────────────────────
    outer = doc.add_table(rows=1, cols=2)
    outer.style = "Table Grid"
    _remove_table_borders(outer)
    from docx.shared import Cm
    outer.columns[0].width = Cm(8.0)
    outer.columns[1].width = Cm(8.0)

    lc = outer.rows[0].cells[0]
    rc = outer.rows[0].cells[1]

    _embed_fin_table(lc, "Exhibit 1: Profit & Loss Statement", is_rows, years)
    _embed_fin_table(rc, "Exhibit 2: Balance Sheet", bs_rows, years)

    doc.add_paragraph()

    # ── CF and Ratios side by side ────────────────────────────────────
    outer2 = doc.add_table(rows=1, cols=2)
    outer2.style = "Table Grid"
    _remove_table_borders(outer2)
    outer2.columns[0].width = Cm(8.0)
    outer2.columns[1].width = Cm(8.0)

    lc2 = outer2.rows[0].cells[0]
    rc2 = outer2.rows[0].cells[1]

    _embed_fin_table(lc2, "Exhibit 3: Cash Flow Statement", cf_rows, years)
    _embed_fin_table(rc2, "Exhibit 4: Key Ratios", rat_rows, years)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"Source: Company, Platform estimates")
    run.font.size = _pt(8)
    run.italic = True
    _color_run(run, _DGRY)


def _embed_fin_table(cell, title: str, rows: List[List[str]], years: List[str]):
    """Write a financial exhibit table inside a cell (for side-by-side layout)."""
    from docx.shared import Pt

    p = cell.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(9)
    _color_run(run, _NAV)
    _shade_paragraph(p, _LBKG)

    if not rows:
        cell.add_paragraph("—")
        return

    ncols = 1 + len(years)
    t = cell.add_table(rows=1 + len(rows), cols=ncols)
    t.style = "Table Grid"

    # Header
    hdr_cells = t.rows[0].cells
    hdr_cells[0].text = "Y/E"
    for ci, yr in enumerate(years, 1):
        hdr_cells[ci].text = str(yr)
    for ci in range(ncols):
        _set_cell_bg(hdr_cells[ci], _TEAL)
        if hdr_cells[ci].paragraphs[0].runs:
            r = hdr_cells[ci].paragraphs[0].runs[0]
            r.bold = True; r.font.size = Pt(7); _color_run(r, _WHT)

    # Data
    for ri, row in enumerate(rows):
        is_bold = row[0].startswith("**") or row[0].isupper()
        bg = _LBKG if is_bold else (_LGY if ri % 2 else _WHT)
        for ci in range(ncols):
            c = t.rows[ri + 1].cells[ci]
            raw = row[ci] if ci < len(row) else _NA
            c.text = str(raw).lstrip("*")
            _set_cell_bg(c, bg)
            if c.paragraphs[0].runs:
                r = c.paragraphs[0].runs[0]
                r.font.size = Pt(7)
                r.bold = is_bold and ci == 0


# ═══════════════════════════════════════════════════════════════════════════
# KEY FINDINGS
# ═══════════════════════════════════════════════════════════════════════════

def _add_findings_section(doc, state):
    _section_bar(doc, "Key Research Findings")

    green_flags = [f for f in state.all_findings if f.finding_type == FindingType.GREEN_FLAG]
    if green_flags:
        _subsection_label(doc, "Positive Indicators")
        for f in green_flags[:8]:
            p = doc.add_paragraph(f"✓  {f.title}", style="List Bullet")
            if p.runs:
                p.runs[0].font.size = _pt(10)
                _color_run(p.runs[0], _GRN)

    risk_text = state.report_sections.get(SectionType.RISK_ANALYSIS.value, "")
    if risk_text:
        _section_bar(doc, "Risk Analysis", color=_LBLU)
        _add_content_paragraphs(doc, risk_text, max_chars=1500)

    scen_text = state.report_sections.get(SectionType.SCENARIO_ANALYSIS.value, "")
    if scen_text:
        _section_bar(doc, "Scenario Narrative", color=_LBLU)
        _add_content_paragraphs(doc, scen_text, max_chars=1000)


# ═══════════════════════════════════════════════════════════════════════════
# CERTIFICATION
# ═══════════════════════════════════════════════════════════════════════════

def _add_certification(doc, company, ticker, report_date, state):
    _section_bar(doc, "Analyst Certification")
    doc.add_paragraph(
        f"The views expressed in this research report accurately reflect the independent "
        f"analysis of {company} ({ticker}) as of {report_date}. No part of the compensation "
        f"of the research team was, is, or will be directly or indirectly related to the "
        f"specific recommendations or views expressed herein. "
        f"Platform: {REPORT_CONFIG.firm_name} v1.0   |   Run ID: {state.run_id}"
    ).runs[0].font.size = _pt(9)


# ═══════════════════════════════════════════════════════════════════════════
# RATING HISTORY  (Deven Choksey — Date / CMP / TP / Recommendation table)
# ═══════════════════════════════════════════════════════════════════════════

def _add_rating_history(doc, company, ticker):
    from docx.shared import Pt

    outer = doc.add_table(rows=1, cols=2)
    outer.style = "Table Grid"
    _remove_table_borders(outer)
    from docx.shared import Cm
    outer.columns[0].width = Cm(8.0)
    outer.columns[1].width = Cm(8.0)

    lc = outer.rows[0].cells[0]
    rc = outer.rows[0].cells[1]

    # Left: rating history
    p = lc.add_paragraph()
    run = p.add_run(f"{ticker}")
    run.bold = True
    run.font.size = Pt(10)
    _color_run(run, _NAV)
    _shade_paragraph(p, _LBKG)

    hist_hdr = ["Date", "CMP (INR)", "TP (INR)", "Recommendation"]
    hist_t = lc.add_table(rows=2, cols=4)
    hist_t.style = "Table Grid"
    for ci, h in enumerate(hist_hdr):
        c = hist_t.rows[0].cells[ci]
        c.text = h
        _set_cell_bg(c, _TEAL)
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].runs[0].font.size = Pt(8)
            _color_run(c.paragraphs[0].runs[0], _WHT)
    # Placeholder current row
    placeholder = [datetime.now().strftime("%d-%b-%y"), "—", "—", "Current Coverage"]
    for ci, v in enumerate(placeholder):
        c = hist_t.rows[1].cells[ci]
        c.text = v
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].font.size = Pt(8)

    # Right: Rating legend
    p2 = rc.add_paragraph()
    run2 = p2.add_run("Rating Legend (Expected over 12 months)")
    run2.bold = True
    run2.font.size = Pt(10)
    _color_run(run2, _NAV)
    _shade_paragraph(p2, _LBKG)

    legend = [
        ("Our Rating", "Upside"),
        ("Buy",        "> 15%"),
        ("Accumulate", "5% – 15%"),
        ("Hold",       "0% – 5%"),
        ("Reduce",     "−5% – 0%"),
        ("Sell",       "< −5%"),
    ]
    leg_t = rc.add_table(rows=len(legend), cols=2)
    leg_t.style = "Table Grid"
    for ri, (k, v) in enumerate(legend):
        bg = _TEAL if ri == 0 else _WHT
        for ci, txt in enumerate((k, v)):
            c = leg_t.rows[ri].cells[ci]
            c.text = txt
            _set_cell_bg(c, bg)
            if c.paragraphs[0].runs:
                r = c.paragraphs[0].runs[0]
                r.font.size = Pt(9)
                r.bold = (ri == 0)
                if ri == 0:
                    _color_run(r, _WHT)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════
# RATING LEGEND (standalone — in main flow after rating history)
# ═══════════════════════════════════════════════════════════════════════════

def _add_rating_legend(doc):
    _section_bar(doc, "Explanation of Investment Rating")
    legend = [
        ("Investment Rating", "Expected Return (over 12 months)"),
        ("STRONG BUY",        "> 25%"),
        ("BUY",               ">= 15%"),
        ("ACCUMULATE",        "5% to 15%"),
        ("HOLD / NEUTRAL",    "−5% to 5%"),
        ("REDUCE",            "−5% to −15%"),
        ("SELL",              "< −15%"),
        ("UNDER REVIEW",      "Rating may undergo a change"),
        ("NOT RATED",         "Forward-looking estimates exist; no recommendation assigned"),
    ]
    t = doc.add_table(rows=len(legend), cols=2)
    t.style = "Table Grid"
    for ri, (k, v) in enumerate(legend):
        bg = _TEAL if ri == 0 else _WHT
        for ci, txt in enumerate((k, v)):
            c = t.rows[ri].cells[ci]
            c.text = txt
            _set_cell_bg(c, bg)
            if c.paragraphs[0].runs:
                r = c.paragraphs[0].runs[0]
                r.font.size = _pt(9)
                r.bold = (ri == 0)
                if ri == 0:
                    _color_run(r, _WHT)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════
# DISCLAIMER
# ═══════════════════════════════════════════════════════════════════════════

def _add_disclaimer(doc, company, report_date):
    _section_bar(doc, "Important Disclosures & Disclaimer")

    blocks = [
        (
            f"RESEARCH ANALYST: Equity Intelligence Agent (AI-assisted research)   |   "
            f"Platform: {REPORT_CONFIG.firm_name}   |   Date: {report_date}"
        ),
        (
            f"This research report has been prepared by {REPORT_CONFIG.firm_name} "
            f"('the Platform') for informational purposes only and does not constitute "
            f"investment advice, a solicitation, or an offer to buy or sell any security."
        ),
        (
            "This report is generated by an AI-assisted autonomous research platform. "
            "All data has been sourced from publicly available information and has not been "
            "independently verified. No representation or warranty, express or implied, is "
            "made as to the accuracy or completeness of the information contained herein."
        ),
        (
            "This report is intended solely for institutional investors and qualified "
            "persons as defined under applicable securities regulations. Past performance "
            "is not indicative of future results. All investments carry risk."
        ),
        (
            "MiFID II Disclosure: This document is a non-independent research communication. "
            "SEBI Disclosure: Produced in compliance with SEBI Research Analyst Regulations (2014) "
            "where applicable. DCFPL or its Research analysts shall be solely responsible for "
            "the security, confidentiality and integrity of the client data."
        ),
        (
            f"The Platform and its associates may have financial interest in {company}. "
            f"Recipients should conduct their own due diligence and seek independent financial "
            f"advice before making any investment decisions."
        ),
        (
            "Registration granted by SEBI, enlistment with Exchange and certification from "
            "NISM in no way guarantee performance of the intermediary or provide any assurance "
            "of returns to investors. Investment in securities market is subject to market risks. "
            "Read all related documents carefully before investing."
        ),
        (
            f"The securities quoted are for illustration only and are not recommendatory. "
            f"© {datetime.now().year} {REPORT_CONFIG.firm_name}. All rights reserved."
        ),
    ]

    for i, block in enumerate(blocks):
        p = doc.add_paragraph(block)
        p.runs[0].font.size = _pt(8)
        if i == 0:
            p.runs[0].bold = True


# ═══════════════════════════════════════════════════════════════════════════
# DATA COLLECTION
# ═══════════════════════════════════════════════════════════════════════════

def _collect_financial_data(state: ResearchState) -> dict:
    history    = state.financial_history or {}
    hist_years = sorted([k for k in history.keys() if len(k) == 4])[-5:]

    fm_out    = state.agent_outputs.get("07_financial_modeling")
    forecasts = (fm_out.payload.get("forecasts", {}) if fm_out else {})
    base_f    = forecasts.get("BASE", [])[:3]

    est_labels = [str(f.get("year", f"FY+{i+1}E")) + "E"
                  for i, f in enumerate(base_f)]
    years = hist_years + est_labels

    def _hv(yr, *keys):
        d = history.get(yr, {})
        for sect in ("income_statements", "balance_sheets", "cash_flows"):
            sd = d.get(sect) or {}
            for k in keys:
                v = sd.get(k)
                if v is not None:
                    return v
        return None

    def _fv(i, key):
        if i < len(base_f):
            return base_f[i].get(key)
        return None

    def _row(label, hist_keys, fcast_key, pct=False, bold=False):
        prefix = "**" if bold else ""
        cells  = [prefix + label]
        for yr in hist_years:
            v = _hv(yr, *hist_keys)
            cells.append(_fmt_num(v, pct))
        for i in range(len(est_labels)):
            v = _fv(i, fcast_key)
            cells.append(_fmt_num(v, pct) if v is not None else _NA)
        return cells

    is_rows = [
        _row("**Sales",            ["revenue"],               "revenue",          bold=True),
        _row("  Change (%)",       ["revenue_growth_pct"],    "revenue_growth_pct", pct=True),
        _row("Employee Cost",      ["employee_cost", "cogs"], "cogs"),
        _row("Other Cost",         ["other_expenses", "opex"],""),
        _row("**EBITDA",           ["ebitda"],                "ebitda",            bold=True),
        _row("  EBITDA Mgn (%)",   ["ebitda_margin_pct"],     "ebitda_margin",     pct=True),
        _row("Depreciation",       ["depreciation", "da"],    "depreciation"),
        _row("**EBIT",             ["ebit"],                  "ebit",              bold=True),
        _row("  EBIT Mgn (%)",     ["ebit_margin_pct"],       "ebit_margin",       pct=True),
        _row("Other income",       ["other_income"],          "other_income"),
        _row("Finance costs",      ["interest_expense", "finance_costs"], ""),
        _row("**Pre-tax Income",   ["pbt", "pre_tax_income"], "pbt",               bold=True),
        _row("Income tax expense", ["tax_expense"],           "tax"),
        _row("**PAT before MI",    ["net_income", "pat"],     "net_income",        bold=True),
        _row("Minority Interest",  ["minority_interest"],     "minority_interest"),
        _row("**PAT after MI",     ["pat_after_mi", "pat"],   "net_income",        bold=True),
        _row("**Diluted EPS",      ["diluted_eps", "eps"],    "eps",               bold=True),
        _row("Shares (Mn)",        ["shares_outstanding"],    ""),
    ]

    bs_rows = [
        _row("**Total Equity",     ["equity", "total_equity", "net_worth"], "equity", bold=True),
        _row("Minority Interest",  ["minority_interest_bs"],  ""),
        _row("Total Borrowings",   ["total_debt", "gross_debt"], "total_debt"),
        _row("**Capital Employed", ["capital_employed"],      "total_assets",      bold=True),
        _row("Gross Block",        ["gross_block", "ppe_gross"], "ppe_gross"),
        _row("**Net Block",        ["net_block", "ppe_net"],  "ppe_net",           bold=True),
        _row("Goodwill",           ["goodwill"],              ""),
        _row("Intangible Assets",  ["intangibles", "intangible_assets"], ""),
        _row("**Total Non-Curr A", ["total_nca", "non_current_assets"], "", bold=True),
        _row("Investments",        ["investments", "current_investments"], ""),
        _row("Trade Receivables",  ["trade_receivables", "accounts_receivable"], ""),
        _row("Cash & Bank",        ["cash_and_equivalents", "cash"], "cash"),
        _row("**Total Curr. Assets",["current_assets", "total_current_assets"], "current_assets", bold=True),
        _row("**Total Curr. Liab.", ["current_liabilities"],  "current_liabilities", bold=True),
        _row("**Total Assets",     ["total_assets"],          "total_assets",      bold=True),
    ]

    cf_rows = [
        _row("**CFFO",             ["cfo", "operating_cash_flow"], "cfo",           bold=True),
        _row("**CFFI",             ["cfi", "investing_cash_flow"],  ""),
        _row("**CFFF",             ["cff", "financing_cash_flow"],  ""),
        _row("Net Inc/Dec in Cash",["net_change_in_cash"],     ""),
        _row("Closing Cash Bal.",  ["cash_and_equivalents", "cash"], "cash"),
    ]

    rat_rows = [
        _row("EBIT Margin (%)",    ["ebit_margin_pct"],       "",        pct=True),
        _row("Tax rate (%)",       ["effective_tax_rate"],    "",        pct=True),
        _row("NPM (%)",            ["net_margin_pct"],        "",        pct=True),
        _row("**RoE (%)",          ["roe_pct"],               "",        bold=True, pct=True),
        _row("RoCE (%)",           ["roce_pct"],              "",        pct=True),
        _row("**P/E (x)",          ["pe_ratio"],              "",        bold=True),
        _row("EV/EBITDA (x)",      ["ev_ebitda"],             ""),
        _row("P/BV (x)",           ["pb_ratio"],              ""),
        _row("**EPS",              ["eps", "basic_eps"],      "eps",     bold=True),
        _row("Book Value/Share",   ["bvps"],                  "bvps"),
        _row("DPS",                ["dps"],                   "dps"),
        _row("Payout (%)",         ["payout_ratio_pct"],      "",        pct=True),
        _row("Div. Yield (%)",     ["dividend_yield_pct"],    "",        pct=True),
        _row("Debtors Days",       ["dso"],                   ""),
        _row("Asset Turnover (x)", ["asset_turnover"],        ""),
    ]

    return {
        "years":            years,
        "income_statement": is_rows,
        "balance_sheet":    bs_rows,
        "cash_flow":        cf_rows,
        "ratios":           rat_rows,
    }


def _add_fin_mini_table(cell, fin_data: dict, currency: str):
    """Compact key financials & valuation mini-table on cover."""
    years = fin_data.get("years", [])[-3:]
    if not years:
        return

    is_rows  = fin_data.get("income_statement", [])
    rat_rows = fin_data.get("ratios", [])

    def _find(rows, fragment):
        for r in rows:
            if fragment.lower() in r[0].lower():
                return r[-len(years):]
        return [_NA] * len(years)

    mini = [
        ("Y/E",           years),
        ("Revenue",       _find(is_rows, "**Sales")),
        ("EBIT Mgn %",    _find(is_rows, "EBIT Mgn")),
        ("PAT",           _find(is_rows, "PAT after MI")),
        ("EPS",           _find(is_rows, "**Diluted EPS")),
        ("EPS Gr. %",     _find(is_rows, "Change")),
        ("BV/Share",      _find(rat_rows, "Book Value")),
        ("RoE %",         _find(rat_rows, "**RoE")),
        ("RoCE %",        _find(rat_rows, "RoCE")),
        ("P/E (x)",       _find(rat_rows, "**P/E")),
        ("EV/EBITDA (x)", _find(rat_rows, "EV/EBITDA")),
    ]

    p = cell.add_paragraph()
    r = p.add_run("Financials & Valuations")
    r.bold = True
    r.font.size = _pt(8)
    _color_run(r, _NAV)
    _shade_paragraph(p, _LBKG)

    t = cell.add_table(rows=len(mini), cols=1 + len(years))
    t.style = "Table Grid"
    for ri, (label, vals) in enumerate(mini):
        bg = _TEAL if ri == 0 else (_LGY if ri % 2 else _WHT)
        t.rows[ri].cells[0].text = label
        _set_cell_bg(t.rows[ri].cells[0], bg)
        if t.rows[ri].cells[0].paragraphs[0].runs:
            rr = t.rows[ri].cells[0].paragraphs[0].runs[0]
            rr.font.size = _pt(7)
            rr.bold = (ri == 0)
            if ri == 0:
                _color_run(rr, _WHT)
        for ci, v in enumerate(vals[:len(years)], 1):
            c = t.rows[ri].cells[ci]
            c.text = str(v)
            _set_cell_bg(c, bg)
            if c.paragraphs[0].runs:
                rr = c.paragraphs[0].runs[0]
                rr.font.size = _pt(7)
                rr.bold = (ri == 0)
                if ri == 0:
                    _color_run(rr, _WHT)


def _add_shareholding_table(cell, state: ResearchState):
    mktdata  = _payload(state, "04_market_data", "market_data", {})
    sholding = mktdata.get("shareholding", {}) or {}
    if not sholding:
        return

    p = cell.add_paragraph()
    r = p.add_run("Share Holding Pattern (%)")
    r.bold = True
    r.font.size = _pt(8)
    _color_run(r, _NAV)
    _shade_paragraph(p, _LBKG)

    promoter = sholding.get("promoter_pct")
    dii      = sholding.get("dii_pct")
    fii      = sholding.get("fii_pct")
    others   = sholding.get("others_pct")

    q_current = sholding.get("quarter_current", "Current")
    q_prev1   = sholding.get("quarter_prev1",   "Q-1")
    q_prev2   = sholding.get("quarter_prev2",   "Q-2")

    hdr = [f"Particulars (%)", q_current, q_prev1, q_prev2]
    rows = [
        ("Promoters", promoter, sholding.get("promoter_pct_q1"), sholding.get("promoter_pct_q2")),
        ("FIIs",      fii,      sholding.get("fii_pct_q1"),      sholding.get("fii_pct_q2")),
        ("DIIs",      dii,      sholding.get("dii_pct_q1"),      sholding.get("dii_pct_q2")),
        ("Others",    others,   sholding.get("others_pct_q1"),   sholding.get("others_pct_q2")),
        ("Total",     100.0,    100.0,                            100.0),
    ]

    t = cell.add_table(rows=1 + len(rows), cols=4)
    t.style = "Table Grid"
    for ci, h in enumerate(hdr):
        c = t.rows[0].cells[ci]
        c.text = h
        _set_cell_bg(c, _LBLU)
        if c.paragraphs[0].runs:
            rr = c.paragraphs[0].runs[0]
            rr.font.size = _pt(7)
            rr.bold = True
            _color_run(rr, _WHT)

    for ri, (k, v1, v2, v3) in enumerate(rows):
        bg = _LGY if ri % 2 else _WHT
        for ci, v in enumerate((k, v1, v2, v3)):
            c = t.rows[ri + 1].cells[ci]
            c.text = f"{v:.1f}" if isinstance(v, (int, float)) else _NA if v is None else str(v)
            _set_cell_bg(c, bg)
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.size = _pt(7)


def _add_kpi_callouts(cell, state, fin_data):
    """Large percentage KPI boxes — Revenue CAGR and PAT CAGR (Deven Choksey style)."""
    fm_out = state.agent_outputs.get("07_financial_modeling")
    if not fm_out:
        return
    base_f = fm_out.payload.get("forecasts", {}).get("BASE", [])
    if len(base_f) < 2:
        return

    def _cagr(key, years=2):
        try:
            v0 = float(base_f[0].get(key, 0))
            v1 = float(base_f[min(years, len(base_f) - 1)].get(key, 0))
            if v0 and v1:
                return ((v1 / v0) ** (1 / years) - 1) * 100
        except Exception:
            pass
        return None

    rev_cagr = _cagr("revenue")
    pat_cagr = _cagr("net_income")
    yrs = len(base_f)

    from docx.shared import Pt
    cell.add_paragraph()

    for label, val in [
        (f"Revenue CAGR\nbetween FY-FY+{yrs}E", rev_cagr),
        (f"Adj. PAT CAGR\nbetween FY-FY+{yrs}E", pat_cagr),
    ]:
        if val is None:
            continue
        p = cell.add_paragraph()
        r_num = p.add_run(f"{val:.1f}%")
        r_num.bold = True
        r_num.font.size = Pt(22)
        _color_run(r_num, _ORG2)
        _shade_paragraph(p, _LBKG)

        p2 = cell.add_paragraph()
        r2 = p2.add_run(label)
        r2.font.size = Pt(8)
        _color_run(r2, _DGRY)


# ═══════════════════════════════════════════════════════════════════════════
# XML / STYLE HELPERS
# ═══════════════════════════════════════════════════════════════════════════

def _pt(size: int):
    from docx.shared import Pt
    return Pt(size)


def _set_margins(doc):
    from docx.shared import Cm
    for sec in doc.sections:
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(1.8)
        sec.right_margin  = Cm(1.8)


def _set_cell_bg(cell, hex_color: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#").upper())
    for ex in tcPr.findall(qn("w:shd")):
        tcPr.remove(ex)
    tcPr.append(shd)


def _shade_paragraph(p, hex_color: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#").upper())
    for ex in pPr.findall(qn("w:shd")):
        pPr.remove(ex)
    pPr.append(shd)


def _color_run(run, hex_color: str):
    from docx.shared import RGBColor
    h = hex_color.lstrip("#")
    run.font.color.rgb = RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _remove_table_borders(table):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl   = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblB  = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblB.append(el)
    for ex in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(ex)
    tblPr.append(tblB)


def _section_bar(doc, text: str, color: str = _TEAL):
    p = doc.add_paragraph()
    _shade_paragraph(p, color)
    run = p.add_run("  " + text)
    run.bold = True
    run.font.size = _pt(10)
    _color_run(run, _WHT)
    return p


def _add_cell_section_bar(cell, text: str, color: str = _TEAL):
    p = cell.add_paragraph()
    _shade_paragraph(p, color)
    run = p.add_run("  " + text)
    run.bold = True
    run.font.size = _pt(9)
    _color_run(run, _WHT)


def _subsection_label(doc, text: str):
    from docx.shared import Pt
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    _color_run(run, _NAV)


def _mini_table(cell, rows: list, hdr_color: str = _LBKG,
                hdr_text: str = "", key_width=None):
    if hdr_text:
        p = cell.add_paragraph()
        r = p.add_run(hdr_text)
        r.bold = True
        r.font.size = _pt(8)
        _color_run(r, _WHT)
        _shade_paragraph(p, _TEAL)

    t = cell.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    for ri, (k, v) in enumerate(rows):
        bg = _LGY if ri % 2 else _WHT
        for ci, txt in enumerate((k, v)):
            c = t.rows[ri].cells[ci]
            c.text = str(txt)
            _set_cell_bg(c, bg)
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.size = _pt(8)
    return t


def _add_bullets_to_cell(cell, items: list, size: int = 9):
    from docx.shared import Pt
    for item in items:
        if item:
            p = cell.add_paragraph()
            run = p.add_run(f"■  {item}")
            run.font.size = Pt(size)


def _add_content_paragraphs(doc, text: str, max_chars: int = 2000):
    if not text:
        return
    text = text[:max_chars]
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith(("##", "**")):
            p = doc.add_paragraph()
            run = p.add_run(block.lstrip("#*").strip())
            run.bold = True
            run.font.size = _pt(10)
        else:
            p = doc.add_paragraph(block)
            if p.runs:
                p.runs[0].font.size = _pt(10)


# ═══════════════════════════════════════════════════════════════════════════
# FORMATTING UTILITIES
# ═══════════════════════════════════════════════════════════════════════════

def _payload(state: ResearchState, agent_id: str, key: str, default=None):
    out = state.agent_outputs.get(agent_id)
    if out is None:
        return default
    return out.payload.get(key, default)


def _fmt(v, pct: bool = False) -> str:
    if v is None:
        return _NA
    try:
        f = float(v)
        if pct:
            return f"{f:.1f}"
        if abs(f) >= 1_000_000:
            return f"{f/1_000_000:.1f}M"
        if abs(f) >= 1_000:
            return f"{f:,.0f}"
        return f"{f:.2f}"
    except (TypeError, ValueError):
        return str(v)


def _fmt_num(v, pct: bool = False) -> str:
    if v is None:
        return _NA
    try:
        f = float(v)
        if pct:
            return f"{f:.1f}"
        if abs(f) >= 1_000:
            return f"{f:,.1f}"
        return f"{f:.1f}"
    except (TypeError, ValueError):
        return str(v)


def _fmt_metric(v, label: str) -> str:
    if v is None:
        return _NA
    pct_hints = ("%", "margin", "growth", "rate", "yield", "payout", "cagr")
    is_pct = any(h in label.lower() for h in pct_hints)
    return _fmt_num(v, pct=is_pct)


def _extract_headline(text: str, company: str) -> str:
    if not text:
        return f"{company}: Equity Research"
    first = text.strip().split("\n")[0].strip()
    return first[:120] if first else f"{company}: Equity Research"


def _extract_bullets(text: str, max_bullets: int = 5) -> List[str]:
    if not text:
        return []
    bullets = []
    for line in text.split("\n"):
        line = line.strip()
        if line.startswith(("- ", "• ", "* ", "■ ")):
            bullets.append(line[2:].strip())
        elif line[:2] in ("1.", "2.", "3.", "4.", "5."):
            bullets.append(line[2:].strip())
        if len(bullets) >= max_bullets:
            break
    if not bullets:
        sentences = [s.strip() for s in text.replace("\n", " ").split(". ") if len(s) > 20]
        bullets = sentences[:max_bullets]
    return [b[:150] for b in bullets if b]
